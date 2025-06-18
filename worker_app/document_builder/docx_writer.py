import os

from typing import List, Optional

from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from worker_app.document_builder.resume_writer_interface import ResumeWriter
from worker_app.document_builder.writer_types import (
    Alignment,
    Font,
    FontStyle,
    TableData,
    TextLine,
)


ALIGNMENT_MAP = {
    Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
    Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
    Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
}

WIDTH = Inches(8.5)


class DOCXWriter(ResumeWriter):
    def __init__(
        self,
        file_path: str,
        font: Font,
        default_font_size=12,
        default_font_style=FontStyle.REGULAR,
        margin_top=72,
        margin_bottom=72,
        margin_left=72,
        margin_right=72,
        background_color="#ffffff",
        text_color="#000000",
    ):
        self.document = Document()
        section = self.document.sections[0]

        background = OxmlElement("w:background")
        background.set(qn("w:color"), background_color.lstrip("#"))

        self.document._element.insert(0, background)

        self.bullet_point_count = 0

        super().__init__(
            font=font,
            file_path=f"{os.environ.get('COMPLETED_RESUMES_DIR', '/app/common/completed_resumes/')}{file_path}.docx",
            default_font_size=default_font_size,
            default_font_style=default_font_style,
            margin_top=margin_top,
            margin_bottom=margin_bottom,
            margin_left=margin_left,
            margin_right=margin_right,
            background_color=background_color,
            text_color=text_color,
        )

        style = self.document.styles["Normal"]
        style.font.color.rgb = RGBColor(
            int(text_color.lstrip("#")[0:2], 16),
            int(text_color.lstrip("#")[2:4], 16),
            int(text_color.lstrip("#")[4:6], 16),
        )

        section.top_margin = Inches(self.margin_top / 96)
        section.bottom_margin = Inches(self.margin_bottom / 96)
        section.left_margin = Inches(self.margin_left / 96)
        section.right_margin = Inches(self.margin_right / 96)

    def save(self):
        self.document.save(self.file_path)
        return self.file_path

    def add_footer_line(
        self,
        value: str | TextLine,
        font_size: int = 12,
        font_style: FontStyle = FontStyle.LIGHT,
        alignment: Alignment = Alignment.LEFT,
    ):
        section = self.document.sections[0]
        footer = section.footer

        paragraph = footer.add_paragraph()
        paragraph.alignment = ALIGNMENT_MAP[alignment]
        if isinstance(value, str):
            run = paragraph.add_run(value)
            font = run.font
            font.name = self._get_font(font_style)
            font.size = Pt(font_size)
        else:
            for part in value.parts:
                run = paragraph.add_run(part["value"])
                font = run.font
                font.name = self._get_font(part["font_style"])
                font.size = Pt(part["font_size"])

    def add_bullet_points(
        self, points: TableData, space_from_point=10, line_spacing=1.5
    ):
        for point in points:
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(point["value"])
            font = run.font
            font.name = self._get_font(point["font_style"])
            font.size = Pt(point["font_size"])

            p = paragraph._p  # access to xml paragraph element
            pPr = p.get_or_add_pPr()  # access paragraph properties

            paragraph.paragraph_format.space_before = Pt(line_spacing)
            paragraph.paragraph_format.space_after = Pt(line_spacing)

            numPr = OxmlElement("w:numPr")  # create number properties element
            numId = OxmlElement("w:numId")  # create numId element - sets bullet type
            numId.set(qn("w:val"), "3")  # set list type/indentation
            numPr.append(numId)  # add bullet type to number properties list

            ilvl = OxmlElement("w:ilvl")  # Not sure
            ilvl.set(qn("w:val"), "0")  # Not sure
            numPr.append(ilvl)

            pPr.append(numPr)

            ind = OxmlElement("w:ind")  # indent
            ind.set(qn("w:left"), "360")
            ind.set(qn("w:hanging"), "360")
            pPr.append(ind)

            rPr = OxmlElement("w:rPr")

            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), self._get_font(point["font_style"]))
            rFonts.set(qn("w:hAnsi"), self._get_font(point["font_style"]))
            rPr.append(rFonts)

            color = OxmlElement("w:color")
            color.set(qn("w:val"), self.text_color.lstrip("#"))
            rPr.append(color)

            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "22")
            rPr.append(sz)

            pPr.append(rPr)  # add bullet point

    def add_table(
        self, data: TableData, column_widths: Optional[List[int]] = None, row_padding=0
    ):
        if not column_widths:
            column_widths = [1] * len(data[0])

        self._validate_table_params(data, column_widths)

        available_width = (
            WIDTH - Inches(self.margin_left / 96) - Inches(self.margin_right / 96)
        )

        total_column_widths = sum(column_widths)

        column_widths_emu = [
            available_width * (width / total_column_widths) for width in column_widths
        ]

        table = self.document.add_table(rows=len(data), cols=len(data[0]))

        for col_index, col in enumerate(table.columns):
            col.width = Emu(column_widths_emu[col_index])

        table.style.font.size = 1

        style = self.document.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(row_padding)

        for row_index, row in enumerate(data):
            table_row = table.rows[row_index]
            max_height = 0
            for col_index, cell in enumerate(row):
                table_cell = table_row.cells[col_index]
                table_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph = table_cell.paragraphs[0]
                paragraph.paragraph_format.left_indent = 0
                paragraph.paragraph_format.line_spacing = 1
                paragraph.alignment = ALIGNMENT_MAP[cell["alignment"]]
                if isinstance(cell["value"], str):
                    run = paragraph.add_run(cell["value"])
                    font = run.font
                    font.name = self._get_font(cell["font_style"])
                    font.size = Pt(cell["font_size"])
                    max_height = max(max_height, cell["font_size"])
                else:
                    for part in cell["value"].parts:
                        run = paragraph.add_run(part["value"])
                        font = run.font
                        font.name = self._get_font(part["font_style"])
                        font.size = Pt(part["font_size"])
                        max_height = max(max_height, part["font_size"])

    def add_vertical_space(self, space: int):
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.line_spacing = Pt(space * 0.75)  # Convert to points

    def write_text_line(self, line: TextLine):
        paragraph = self.document.add_paragraph()
        for part in line.parts:
            run = paragraph.add_run(part["value"])
            font = run.font
            font.name = self._get_font(part["font_style"])
            font.size = Pt(part["font_size"])
            max_height = max(max_height, part["font_size"])

    def add_text(
        self,
        value: str,
        font_size: int,
        font_style: FontStyle = FontStyle.SEMIBOLD,
        alignment: Alignment = Alignment.CENTER,
        line_spacing: float = 1.0,
    ):
        paragraph = self.document.add_paragraph()
        paragraph.alignment = ALIGNMENT_MAP[alignment]
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.space_before = Pt(1)
        run = paragraph.add_run(value)
        font = run.font
        font.name = self._get_font(font_style)
        font.size = Pt(font_size)

    def add_horizontal_line(self, width=1):
        paragraph = self.document.add_paragraph()
        paragraph.style.paragraph_format.space_before = Pt(1)
        paragraph.style.paragraph_format.space_after = Pt(8)
        p = paragraph._p  # Access the XML of the paragraph

        p.clear_content()

        pPr = OxmlElement("w:pPr")

        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:space"), "1")  # space above line
        bottom.set(qn("w:sz"), f"{width * 7}")  # width of line
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:color"), self.text_color.lstrip("#"))
        pBdr.append(bottom)
        pPr.append(pBdr)

        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "240")  # unknown
        spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)

        rPr = OxmlElement("w:rPr")

        b = OxmlElement("w:b")
        b.set(qn("w:val"), "1")
        rPr.append(b)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "ffffff")
        rPr.append(color)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "1")  # space above line
        rPr.append(sz)

        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "1")  # space above line
        rPr.append(szCs)

        pPr.append(rPr)

        p.append(pPr)
